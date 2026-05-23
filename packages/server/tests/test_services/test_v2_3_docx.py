"""Tests for the v2.3 docx renderer.

Verifies the structural contract — valid Word document, cover, sections,
footnotes section, chart data table — without snapshotting the full
binary (which would be brittle across python-docx versions).
"""

from __future__ import annotations

import io
import zipfile
from datetime import UTC, datetime

import pytest
from docx import Document
from openlia.llm.runtime.report_v2_3.schemas import (
    BundleFact,
    BundleSeries,
    BundleSeriesPoint,
    CanonicalFigure,
    ChartSeries,
    ChartSpec,
    ChartType,
    DataProviderSource,
    Language,
    Outline,
    OutlineSection,
    ReportThesis,
    ReportType,
    ResearchBundle,
    SectionMandate,
    ValuationPlan,
    WrittenSection,
)
from openlia.llm.runtime.report_v2_3.stages import StageContext
from openlia.llm.runtime.report_v2_3.stages.assemble import AssembleStage
from openlia.llm.runtime.report_v2_3.state import ReportState
from openlia_server.services.v2_3_docx import render_docx


def _src() -> DataProviderSource:
    return DataProviderSource(
        provider="EODHD",
        endpoint="fundamentals/income_statement",
        period="TTM",
        retrieved_at=datetime.now(UTC),
    )


def _bundle() -> ResearchBundle:
    return ResearchBundle(
        tickers=["NVDA"],
        facts={
            "rev_series": BundleFact(
                id="rev_series",
                label="Revenue (quarterly)",
                value=BundleSeries(
                    points=[
                        BundleSeriesPoint(period="2025-Q1", value=10.0),
                        BundleSeriesPoint(period="2025-Q2", value=11.5),
                    ],
                    unit="USD_billions",
                ),
                source=_src(),
            ),
            "gm": BundleFact(id="gm", label="Gross margin", value=0.65, source=_src()),
        },
    )


def _thesis_with_chart() -> ReportThesis:
    return ReportThesis(
        language=Language.EN,
        central_argument="Durable growth in data-center revenue.",
        key_takeaways=["DC revenue scaled", "Gross margins held"],
        valuation_stance="Fair value above current price.",
        valuation_plan=ValuationPlan(),
        canonical_figures=[CanonicalFigure(fact_id="gm", display="65.0%")],
        mandates=[
            SectionMandate(
                section_id="overview",
                covers="business overview",
                does_not_cover="financials",
                relevant_fact_ids=["gm"],
            ),
            SectionMandate(
                section_id="financials",
                covers="financials",
                does_not_cover="overview",
                chart_ids=["rev_chart"],
                relevant_fact_ids=["rev_series", "gm"],
            ),
        ],
        charts=[
            ChartSpec(
                id="rev_chart",
                section_id="financials",
                claim="revenue scaled meaningfully in 2025",
                chart_type=ChartType.LINE,
                title="Revenue (quarterly)",
                category_labels=["2025-Q1", "2025-Q2"],
                x_axis_label="Quarter",
                series=[ChartSeries(name="revenue", value_fact_ids=["rev_series"])],
            )
        ],
    )


def _state_ready_to_render() -> ReportState:
    s = ReportState(
        run_id="r",
        user_id="u",
        raw_prompt="initiate on NVDA",
        language=Language.EN,
        report_type=ReportType.INITIATION,
        tickers=["NVDA"],
    )
    s.bundle = _bundle()
    s.outline = Outline(
        tickers=["NVDA"],
        report_type=ReportType.INITIATION,
        sections=[
            OutlineSection(id="overview", title="Overview"),
            OutlineSection(id="financials", title="Financials"),
        ],
    )
    s.thesis = _thesis_with_chart()
    s.sections = [
        WrittenSection(
            section_id="overview",
            title="Overview",
            body="Margins stayed at {{CITE:gm}} despite competitive pressure.",
        ),
        WrittenSection(
            section_id="financials",
            title="Financials",
            body="Quarterly revenue rose materially; see {{FIG:rev_chart}}.",
        ),
    ]
    # AssembleStage runs first to populate state.resolved.
    AssembleStage().run(s, StageContext(clients={}, tools={}, extras={}))
    return s


# ---------------------------------------------------------------------------
# Renderer output shape
# ---------------------------------------------------------------------------


def test_render_returns_valid_docx_zip() -> None:
    blob = render_docx(_state_ready_to_render())
    assert isinstance(blob, bytes) and len(blob) > 0
    assert zipfile.is_zipfile(io.BytesIO(blob))


def test_rendered_document_contains_cover_and_section_titles() -> None:
    blob = render_docx(_state_ready_to_render())
    doc = Document(io.BytesIO(blob))
    # Cover uses level=0 ("Title" style); section/thesis headings use Heading 1/2.
    titled = [
        p.text
        for p in doc.paragraphs
        if p.style.name.startswith("Heading") or p.style.name == "Title"
    ]
    assert "NVDA" in titled
    assert "Overview" in titled
    assert "Financials" in titled
    assert "Thesis" in titled


def test_rendered_document_resolves_footnote_markers() -> None:
    blob = render_docx(_state_ready_to_render())
    doc = Document(io.BytesIO(blob))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    # PR16: ``[^1]`` was rewritten into a real Word footnote reference so it
    # no longer appears as visible text; ``{{CITE:}}`` likewise.
    assert "[^1]" not in body_text
    assert "{{CITE:" not in body_text
    # Figure was numbered and the placeholder was replaced.
    assert "Figure 1" in body_text
    assert "{{FIG:" not in body_text
    # Native footnote reference and the footnotes part both exist.
    body_xml = doc.part.element.body.xml
    assert "footnoteReference" in body_xml
    assert any(rel.reltype.endswith("/footnotes") for rel in doc.part.rels.values())


def test_rendered_document_contains_references_section() -> None:
    blob = render_docx(_state_ready_to_render())
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    assert "References" in paragraphs
    # Each footnote becomes a "[N] ..." line listing the provider call.
    assert any(p.startswith("[1] EODHD") for p in paragraphs)


def test_rendered_document_contains_chart_data_table() -> None:
    blob = render_docx(_state_ready_to_render())
    doc = Document(io.BytesIO(blob))
    assert len(doc.tables) >= 1
    # Find a table whose first cell matches our x_axis_label.
    revenue_table = next(t for t in doc.tables if t.rows[0].cells[0].text == "Quarter")
    # Header row: ["Quarter", "revenue"].
    assert revenue_table.rows[0].cells[1].text == "revenue"
    # Body rows mirror the BundleSeries values, not the fact_id.
    assert revenue_table.rows[1].cells[0].text == "2025-Q1"
    assert revenue_table.rows[1].cells[1].text == "10.00"
    assert revenue_table.rows[2].cells[0].text == "2025-Q2"
    assert revenue_table.rows[2].cells[1].text == "11.50"


# ---------------------------------------------------------------------------
# Precondition errors
# ---------------------------------------------------------------------------


def test_render_raises_when_resolved_missing() -> None:
    state = _state_ready_to_render()
    state.resolved = None
    with pytest.raises(RuntimeError, match=r"state\.resolved"):
        render_docx(state)


def test_render_raises_when_thesis_missing() -> None:
    state = _state_ready_to_render()
    state.thesis = None
    with pytest.raises(RuntimeError, match="outline"):
        render_docx(state)
