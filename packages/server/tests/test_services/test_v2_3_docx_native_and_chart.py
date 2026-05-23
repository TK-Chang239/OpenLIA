"""Targeted tests for the v2.3 docx native-footnote + chart-PNG modules."""

from __future__ import annotations

import io
from datetime import UTC, datetime

from docx import Document
from openlia.llm.runtime.report_v2_3.schemas import (
    BundleFact,
    BundleSeries,
    BundleSeriesPoint,
    ChartSeries,
    ChartSpec,
    ChartType,
    DataProviderSource,
    ResearchBundle,
)
from openlia_server.services.v2_3_docx_chart import try_render_chart_png
from openlia_server.services.v2_3_docx_native import attach_native_footnotes


def _src() -> DataProviderSource:
    return DataProviderSource(
        provider="EODHD",
        endpoint="fundamentals",
        period="TTM",
        retrieved_at=datetime.now(UTC),
    )


# ---------------------------------------------------------------------------
# Native footnotes
# ---------------------------------------------------------------------------


def test_attach_native_footnotes_adds_footnotes_part_and_rewrites_markers() -> None:
    doc = Document()
    doc.add_paragraph("Revenue grew sharply [^1] in 2025 [^2].")
    attach_native_footnotes(doc, ["First footnote text.", "Second footnote text."])

    # Markers are no longer visible text.
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[^1]" not in body_text
    assert "[^2]" not in body_text

    # Surrounding prose survived.
    assert "Revenue grew sharply" in body_text
    assert "in 2025" in body_text

    # Native footnote references in the body XML.
    body_xml = doc.part.element.body.xml
    assert body_xml.count("footnoteReference") == 2

    # Footnotes part attached with the expected relationship type.
    rels = doc.part.rels
    footnote_rels = [r for r in rels.values() if r.reltype.endswith("/footnotes")]
    assert len(footnote_rels) == 1
    blob = footnote_rels[0].target_part.blob
    assert b"First footnote text." in blob
    assert b"Second footnote text." in blob
    # The standard separator entries are present so Word can render the
    # horizontal line above the footnote block.
    assert b'w:type="separator"' in blob
    assert b'w:type="continuationSeparator"' in blob


def test_attach_native_footnotes_no_op_when_empty() -> None:
    doc = Document()
    doc.add_paragraph("No citations here.")
    attach_native_footnotes(doc, [])
    rels = doc.part.rels
    assert not any(r.reltype.endswith("/footnotes") for r in rels.values())


def test_attach_native_footnotes_preserves_run_properties() -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("bold prefix [^1] bold suffix")
    run.bold = True

    attach_native_footnotes(doc, ["Footnote."])

    # The split should preserve the bold formatting on the prose halves.
    body_xml = doc.part.element.body.xml
    assert "<w:b/>" in body_xml or 'w:val="true"' in body_xml or "w:b " in body_xml


def test_attach_native_footnotes_writes_valid_docx() -> None:
    doc = Document()
    doc.add_paragraph("Hello [^1] world.")
    attach_native_footnotes(doc, ["Footnote."])

    buf = io.BytesIO()
    doc.save(buf)
    # Re-open round trip — corrupted docx would raise.
    reopened = Document(io.BytesIO(buf.getvalue()))
    assert "Hello" in reopened.paragraphs[0].text


# ---------------------------------------------------------------------------
# Chart PNG rendering
# ---------------------------------------------------------------------------


def _bundle_with_series(n: int = 3) -> ResearchBundle:
    return ResearchBundle(
        tickers=["NVDA"],
        facts={
            "rev_series": BundleFact(
                id="rev_series",
                label="Revenue series",
                value=BundleSeries(
                    points=[
                        BundleSeriesPoint(period=f"FY{2020 + i}", value=float(50 + i * 10))
                        for i in range(n)
                    ],
                    unit="USD_billions",
                ),
                source=_src(),
            ),
        },
    )


def _bundle_with_scalars() -> ResearchBundle:
    return ResearchBundle(
        tickers=["NVDA"],
        facts={
            "a": BundleFact(id="a", label="A", value=10.0, source=_src()),
            "b": BundleFact(id="b", label="B", value=20.0, source=_src()),
            "c": BundleFact(id="c", label="C", value=30.0, source=_src()),
        },
    )


def _chart(chart_type: ChartType, *, series_layout: str) -> ChartSpec:
    if series_layout == "bundle_series":
        series = [ChartSeries(name="Revenue", value_fact_ids=["rev_series"])]
    else:  # one_fact_per_category
        series = [ChartSeries(name="Group", value_fact_ids=["a", "b", "c"])]
    return ChartSpec(
        id="chart_demo",
        section_id="financials",
        claim="Demo claim",
        chart_type=chart_type,
        title="Demo chart",
        category_labels=["FY2020", "FY2021", "FY2022"],
        series=series,
    )


def test_chart_png_column_bundleseries_layout() -> None:
    bundle = _bundle_with_series(3)
    chart = _chart(ChartType.COLUMN, series_layout="bundle_series")
    png = try_render_chart_png(chart, bundle)
    assert png is not None
    assert png.startswith(b"\x89PNG\r\n\x1a\n")  # PNG magic number


def test_chart_png_line_one_fact_per_category() -> None:
    bundle = _bundle_with_scalars()
    chart = _chart(ChartType.LINE, series_layout="one_fact_per_category")
    png = try_render_chart_png(chart, bundle)
    assert png is not None
    assert png.startswith(b"\x89PNG")


def test_chart_png_pie_renders() -> None:
    bundle = _bundle_with_scalars()
    chart = _chart(ChartType.PIE, series_layout="one_fact_per_category")
    png = try_render_chart_png(chart, bundle)
    assert png is not None


def test_chart_png_returns_none_for_mismatched_lengths() -> None:
    # 3 categories but only 2 points -> cannot render.
    bundle = _bundle_with_series(2)
    chart = _chart(ChartType.COLUMN, series_layout="bundle_series")
    assert try_render_chart_png(chart, bundle) is None


def test_chart_png_returns_none_for_missing_fact() -> None:
    bundle = ResearchBundle(
        tickers=["NVDA"],
        facts={"a": BundleFact(id="a", label="A", value=10.0, source=_src())},
    )
    chart = _chart(ChartType.COLUMN, series_layout="one_fact_per_category")
    # missing 'b' and 'c'
    assert try_render_chart_png(chart, bundle) is None


def test_chart_png_multi_series_bar() -> None:
    bundle = ResearchBundle(
        tickers=["NVDA"],
        facts={
            "a1": BundleFact(id="a1", label="A1", value=10.0, source=_src()),
            "a2": BundleFact(id="a2", label="A2", value=20.0, source=_src()),
            "b1": BundleFact(id="b1", label="B1", value=15.0, source=_src()),
            "b2": BundleFact(id="b2", label="B2", value=25.0, source=_src()),
        },
    )
    chart = ChartSpec(
        id="c",
        section_id="s",
        claim="x",
        chart_type=ChartType.BAR,
        title="t",
        category_labels=["A", "B"],
        series=[
            ChartSeries(name="2024", value_fact_ids=["a1", "b1"]),
            ChartSeries(name="2025", value_fact_ids=["a2", "b2"]),
        ],
    )
    png = try_render_chart_png(chart, bundle)
    assert png is not None
