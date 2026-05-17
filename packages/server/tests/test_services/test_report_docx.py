from __future__ import annotations

import io


def _png_bytes() -> bytes:
    # Pillow ships with python-docx; use it to build a tiny valid PNG.
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (4, 4), color=(255, 0, 0)).save(buf, format="PNG")
    return buf.getvalue()


def _load_doc(buf: bytes):
    from docx import Document

    return Document(io.BytesIO(buf))


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_assemble_docx_includes_cover_and_section_heading() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {
            "title": "Acme — Initiation",
            "subtitle": "Buy at $100",
            "ticker": "ACME",
            "tagline": "Strong moat",
            "key_metrics": [
                {"label": "Price", "value": "$95"},
                {"label": "Target", "value": "$120"},
            ],
        },
        "sections": [
            {
                "id": "thesis",
                "title": "Thesis",
                "blocks": [{"type": "text", "content": "Hello world"}],
            }
        ],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={}, header_text="Acme")
    doc = _load_doc(docx_bytes)
    text = _all_text(doc)
    assert "Acme — Initiation" in text
    assert "Thesis" in text
    assert "Hello world" in text
    assert "Price" in text and "$95" in text


def test_assemble_docx_embeds_chart_png_at_block_path() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [
            {
                "id": "thesis",
                "title": "Thesis",
                "blocks": [
                    {"type": "text", "content": "Intro"},
                    {"type": "pie_chart", "title": "Revenue mix", "segments": []},
                ],
            }
        ],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={"0-1": _png_bytes()}, header_text="Acme")
    doc = _load_doc(docx_bytes)
    image_parts = [p for p in doc.part.related_parts.values() if "image" in p.content_type]
    assert len(image_parts) >= 1
    # the chart's title still renders
    assert "Revenue mix" in _all_text(doc)


def test_assemble_docx_chart_without_png_falls_back_to_placeholder() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [
            {
                "id": "thesis",
                "title": "Thesis",
                "blocks": [{"type": "pie_chart", "title": "Mix", "segments": []}],
            }
        ],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={}, header_text="Acme")
    doc = _load_doc(docx_bytes)
    text = _all_text(doc)
    # Placeholder uses the chart title or type
    assert "Mix" in text


def test_assemble_docx_renders_table_block_as_word_table() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [
            {
                "id": "fin",
                "title": "Financials",
                "blocks": [
                    {
                        "type": "table",
                        "title": "Revenue",
                        "headers": [
                            {"key": "year", "label": "Year"},
                            {"key": "rev", "label": "Revenue"},
                        ],
                        "rows": [
                            {"year": "2024", "rev": "$100M"},
                            {"year": "2025", "rev": "$120M"},
                        ],
                    }
                ],
            }
        ],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={}, header_text="Acme")
    doc = _load_doc(docx_bytes)
    # First table is the metric-cards cover; the section's revenue table is
    # the next one. We just verify cell content rather than table count.
    text = _all_text(doc)
    assert "Year" in text and "Revenue" in text
    assert "2024" in text and "$100M" in text
    assert "2025" in text and "$120M" in text


def test_assemble_docx_renders_metric_cards_as_table() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [
            {
                "id": "kpis",
                "title": "KPIs",
                "blocks": [
                    {
                        "type": "metric_cards",
                        "metrics": [
                            {"label": "Users", "value": "1.2M"},
                            {"label": "MRR", "value": "$3.4M", "delta": "+12%"},
                        ],
                    }
                ],
            }
        ],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={}, header_text="Acme")
    text = _all_text(_load_doc(docx_bytes))
    assert "Users" in text and "1.2M" in text
    assert "MRR" in text and "$3.4M" in text and "+12%" in text


def test_assemble_docx_bullet_list_and_quote() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [
            {
                "id": "s",
                "title": "Section",
                "blocks": [
                    {"type": "bullet_list", "items": ["alpha", "beta"]},
                    {
                        "type": "quote",
                        "content": "Stay hungry",
                        "attribution": "Steve",
                    },
                ],
            }
        ],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={}, header_text="Acme")
    text = _all_text(_load_doc(docx_bytes))
    assert "alpha" in text and "beta" in text
    assert "Stay hungry" in text and "Steve" in text


def test_assemble_docx_group_recurses_into_children() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [
            {
                "id": "s",
                "title": "S",
                "blocks": [
                    {
                        "type": "group",
                        "blocks": [
                            {"type": "text", "content": "child-a"},
                            {"type": "text", "content": "child-b"},
                        ],
                    }
                ],
            }
        ],
    }
    text = _all_text(_load_doc(assemble_docx(schema, chart_pngs={}, header_text="X")))
    assert "child-a" in text and "child-b" in text


def test_assemble_docx_inserts_word_toc_field() -> None:
    from openlia_server.services.report_docx import assemble_docx

    schema = {
        "cover": {"title": "Acme"},
        "sections": [{"id": "a", "title": "A", "blocks": [{"type": "text", "content": "x"}]}],
    }
    docx_bytes = assemble_docx(schema, chart_pngs={}, header_text="Acme")
    # Inspect the document.xml for a TOC field instruction.
    import zipfile

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        body = z.read("word/document.xml").decode("utf-8")
    assert "TOC" in body
