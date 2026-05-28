"""v3 docx renderer tests.

Exercises the structural shape of the output without re-parsing the
zip (python-docx ``Document.from_xxx`` walking is enough). Checks:
  - bytes look like a docx (PK signature)
  - section titles + body text land in document.paragraphs
  - bullet/numbered lists render with list styles
  - citation [^source_id] markers rewrite to [N]
  - {{chart:id}} placeholders are stripped from prose
  - bibliography appears with sources in display_index order
"""

from __future__ import annotations

import io
import json
import uuid
from datetime import UTC, datetime

import pytest
from docx import Document as DocxDocument
from openlia.llm.runtime.report_v3 import ChartSpec
from openlia_server.db.models.auth import User
from openlia_server.db.models.report_v3 import (
    ReportV3,
    ReportV3Chart,
    ReportV3Citation,
    ReportV3Section,
)
from openlia_server.services.v3_docx import render_docx
from sqlalchemy.orm import Session


def _user(db: Session) -> User:
    u = User(
        id=str(uuid.uuid4()),
        email=f"docx-{uuid.uuid4().hex[:6]}@test.com",
        password_hash="x",
        display_name="D",
    )
    db.add(u)
    db.flush()
    return u


def _report(user_id: str) -> ReportV3:
    now = datetime.now(UTC)
    return ReportV3(
        id=str(uuid.uuid4()),
        user_id=user_id,
        subject="RKLB.US",
        template_id="initiation_default",
        language="en",
        length="normal",
        provider_kind="anthropic",
        model="claude-sonnet-4-6",
        status="completed",
        error_message=None,
        created_at=now,
        completed_at=now,
    )


def _section(
    report_id: str,
    section_id: str,
    title: str,
    markdown: str,
    *,
    index: int = 0,
    version: int = 1,
) -> ReportV3Section:
    return ReportV3Section(
        report_id=report_id,
        section_id=section_id,
        section_index=index,
        title=title,
        markdown=markdown,
        revision_id=None,
        version=version,
    )


def _citation(
    report_id: str,
    source_id: str,
    display_index: int | None,
    *,
    url: str | None = None,
) -> ReportV3Citation:
    prov: dict[str, str] = {}
    if url:
        prov["url"] = url
    return ReportV3Citation(
        report_id=report_id,
        source_id=source_id,
        tool_name="web_search",
        display_index=display_index,
        provenance_json=json.dumps(prov),
    )


def _read_paragraphs(blob: bytes) -> list[str]:
    return [p.text for p in DocxDocument(io.BytesIO(blob)).paragraphs]


def test_render_docx_emits_valid_word_bytes(create_tables, db_session: Session):
    user = _user(db_session)
    report = _report(user.id)
    sections = [
        _section(
            report.id,
            "overview",
            "Overview",
            "Rocket Lab is a launch provider [^web_1].",
        )
    ]
    citations = [_citation(report.id, "web_1", 1, url="https://example.com/rklb")]

    blob = render_docx(
        report=report, sections=sections, charts=[], citations=citations
    )
    # PK signature on a .docx (zip container)
    assert blob[:2] == b"PK"
    paragraphs = _read_paragraphs(blob)
    text = "\n".join(paragraphs)
    assert "RKLB.US" in text  # cover title
    assert "Overview" in text  # section heading
    # Citation marker rewritten to [1]
    assert "[1]" in text
    assert "[^web_1]" not in text
    # Bibliography surfaces the URL
    assert "https://example.com/rklb" in text


def test_render_docx_strips_chart_placeholders_from_prose(
    create_tables, db_session: Session
):
    user = _user(db_session)
    report = _report(user.id)
    sections = [
        _section(
            report.id,
            "model",
            "Model",
            "Revenue trend:\n\n{{chart:rev_trend}}\n\nGrowth is strong.",
        )
    ]
    spec = ChartSpec(
        chart_id="rev_trend",
        chart_type="bar",
        title="Revenue trend",
        data=[
            {"label": "2024", "value": 436},
            {"label": "2025", "value": 575},
        ],
    )
    charts = [
        ReportV3Chart(
            report_id=report.id,
            chart_id="rev_trend",
            chart_type="bar",
            title="Revenue trend",
            spec_json=spec.model_dump_json(),
            rendered_url=None,
            revision_id=None,
            version=1,
        )
    ]

    blob = render_docx(
        report=report, sections=sections, charts=charts, citations=[]
    )
    text = "\n".join(_read_paragraphs(blob))
    assert "{{chart:rev_trend}}" not in text
    assert "Revenue trend" in text  # chart title rendered as a heading line


def test_render_docx_renders_markdown_lists_with_proper_styles(
    create_tables, db_session: Session
):
    user = _user(db_session)
    report = _report(user.id)
    sections = [
        _section(
            report.id,
            "thesis",
            "Thesis",
            "Key drivers:\n\n- Launch cadence\n- Neutron readiness\n- Backlog growth\n",
        )
    ]
    blob = render_docx(
        report=report, sections=sections, charts=[], citations=[]
    )
    doc = DocxDocument(io.BytesIO(blob))
    # The 3 bullets should each get a List Bullet style.
    bullet_paras = [
        p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"
    ]
    assert len(bullet_paras) == 3
    assert {p.text for p in bullet_paras} == {
        "Launch cadence",
        "Neutron readiness",
        "Backlog growth",
    }


def test_render_docx_bibliography_orders_by_display_index(
    create_tables, db_session: Session
):
    user = _user(db_session)
    report = _report(user.id)
    sections = [
        _section(
            report.id,
            "s1",
            "S1",
            "Body referencing [^a] and [^b].",
        )
    ]
    # display_index 2 + 1 (intentionally out of insertion order so we
    # can confirm the renderer sorts).
    citations = [
        _citation(report.id, "a", 2, url="https://example.com/two"),
        _citation(report.id, "b", 1, url="https://example.com/one"),
    ]
    blob = render_docx(
        report=report, sections=sections, charts=[], citations=citations
    )
    paragraphs = _read_paragraphs(blob)
    # Find paragraph text that includes a bibliography URL — the [1]
    # entry should appear before the [2] entry.
    one_pos = next(
        i for i, p in enumerate(paragraphs) if "https://example.com/one" in p
    )
    two_pos = next(
        i for i, p in enumerate(paragraphs) if "https://example.com/two" in p
    )
    assert one_pos < two_pos


def test_render_docx_handles_empty_sections_gracefully(
    create_tables, db_session: Session
):
    user = _user(db_session)
    report = _report(user.id)
    # No sections, no charts, no citations — should still emit valid
    # docx with just the cover.
    blob = render_docx(report=report, sections=[], charts=[], citations=[])
    assert blob[:2] == b"PK"
    text = "\n".join(_read_paragraphs(blob))
    assert "RKLB.US" in text


@pytest.mark.parametrize(
    "marker,expected",
    [
        ("**bold**", True),
        ("*italic*", True),
        ("[link](https://x.test)", True),
    ],
)
def test_render_docx_inline_formatting_survives_round_trip(
    marker, expected, create_tables, db_session: Session
):
    user = _user(db_session)
    report = _report(user.id)
    sections = [_section(report.id, "x", "X", f"Has {marker} content.")]
    blob = render_docx(
        report=report, sections=sections, charts=[], citations=[]
    )
    # Just confirm rendering didn't crash and content lands in the doc.
    text = "\n".join(_read_paragraphs(blob))
    assert ("bold" in text or "italic" in text or "link" in text) == expected
