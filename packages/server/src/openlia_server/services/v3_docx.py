"""DOCX renderer for v3 equity-research reports.

Consumes the same persisted rows the HTML / PDF renderers do
(``ReportV3`` + sections + charts + citations) and returns
``.docx`` bytes. Charts try the existing
``report_v3.rendering.chart_renderer`` (matplotlib PNG) and embed
the image; when rendering is impossible (no matplotlib in the env,
mis-shaped data, unsupported chart type) the renderer falls back
to a data table so the figure is never silently dropped.

Markdown parsing uses ``markdown-it-py`` (already a project dep) so
headings / paragraphs / bold / italic / lists / links / code spans
all render natively in Word. ``[^source_id]`` citation markers are
rewritten to ``[N]`` numeric references using each citation's
``display_index`` — native Word footnotes are a future enhancement;
the bracketed numbers match the HTML renderer's output.

The bibliography lands as a "Sources" tail section so readers can
chase citations regardless of viewer capabilities.
"""

from __future__ import annotations

import io
import json
import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from openlia.llm.runtime.report_v3 import ChartSpec, CoverSpec
from openlia.llm.runtime.report_v3.rendering import strip_anthropic_citation_markup

from openlia_server.db.models.report_v3 import (
    ReportV3,
    ReportV3Chart,
    ReportV3Citation,
    ReportV3Section,
)

log = logging.getLogger(__name__)

# Heading levels recognised inside section bodies. The section title
# itself uses level 2 (the report's cover title is level 1), so body
# headings step down from level 3.
_BODY_HEADING_BASE_LEVEL = 3
_MAX_HEADING_LEVEL = 9  # python-docx caps at Heading 9


def render_docx(
    *,
    report: ReportV3,
    sections: list[ReportV3Section],
    charts: list[ReportV3Chart],
    citations: list[ReportV3Citation],
) -> bytes:
    """Build a Word document from the persisted v3 rows."""
    doc = Document()
    _configure_styles(doc)
    cover_spec = _parse_cover(report.cover_json)
    _add_cover(doc, report, cover_spec)

    display_index_by_source_id = {
        c.source_id: c.display_index for c in citations if c.display_index is not None
    }
    charts_by_id = {c.chart_id: c for c in charts}

    md = MarkdownIt("commonmark", {"breaks": True, "html": False}).enable("table")
    for section in sections:
        _add_section(
            doc=doc,
            section=section,
            md=md,
            display_index_by_source_id=display_index_by_source_id,
            charts_by_id=charts_by_id,
        )

    _add_bibliography(doc, citations)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------


def _configure_styles(doc: Document) -> None:
    """Set the default body font to a clean sans-serif at 11pt.
    Headings inherit from the same base via python-docx's built-in
    styles; no per-heading overrides needed for the v3 docx."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


_TONE_COLORS: dict[str, RGBColor] = {
    "positive": RGBColor(0x1A, 0x7F, 0x37),  # green
    "negative": RGBColor(0xB4, 0x23, 0x18),  # red
    "neutral": RGBColor(0x60, 0x60, 0x60),  # gray
}


def _parse_cover(raw: str | None) -> CoverSpec | None:
    """Deserialise ``ReportV3.cover_json`` to a ``CoverSpec``. Returns
    None for runs where the model never called ``set_cover`` (older
    rows + early v3 runs) or when the stored JSON can't be parsed —
    the cover degrades to the bare title + eyebrow path."""
    if not raw:
        return None
    try:
        return CoverSpec.model_validate_json(raw)
    except (TypeError, ValueError):
        log.warning("v3 cover_json failed to deserialise; rendering bare cover")
        return None


def _add_cover(doc: Document, report: ReportV3, cover: CoverSpec | None) -> None:
    """Emit the cover hero block.

    Bare path (cover is None): centered title + eyebrow line, then a
    spacer — the original behaviour. Populated path: adds an optional
    subtitle above the title, a tagline beneath the eyebrow, a
    rating + upside line, a Highlights bullet list, and a key metrics
    table. Empty fields on a populated CoverSpec just skip their row
    rather than render blank elements.
    """
    if cover is not None and cover.subtitle:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_run = subtitle.add_run(cover.subtitle)
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        sub_run.italic = True

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(report.subject)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    eyebrow_run = eyebrow.add_run(_format_eyebrow(report))
    eyebrow_run.font.size = Pt(10)
    eyebrow_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if cover is None:
        doc.add_paragraph()  # spacer before sections
        return

    if cover.tagline:
        tagline = doc.add_paragraph()
        tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tag_run = tagline.add_run(cover.tagline)
        tag_run.italic = True
        tag_run.font.size = Pt(13)
        tag_run.font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    rating_line = _format_rating_line(cover)
    if rating_line:
        rating_para = doc.add_paragraph()
        rating_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        rating_run = rating_para.add_run(rating_line)
        rating_run.bold = True
        rating_run.font.size = Pt(11)
        rating_run.font.color.rgb = RGBColor(0x1A, 0x7F, 0x37)

    if cover.tldr:
        doc.add_paragraph()  # blank line before the Highlights heading
        highlights = doc.add_heading("Highlights", level=3)
        highlights.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for bullet in cover.tldr:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(bullet)

    if cover.key_metrics:
        if not cover.tldr:
            doc.add_paragraph()  # blank line before the metrics table
        _add_key_metrics_table(doc, cover.key_metrics)

    doc.add_paragraph()  # spacer before sections


def _format_rating_line(cover: CoverSpec) -> str:
    """Compose a one-line rating summary, e.g.
    ``"Buy  ·  +28.5% upside"`` or just ``"Hold"`` when no upside.
    Returns empty when neither field is set."""
    parts: list[str] = []
    if cover.rating:
        parts.append(cover.rating)
    if cover.upside_pct is not None:
        sign = "+" if cover.upside_pct >= 0 else ""
        parts.append(f"{sign}{cover.upside_pct:g}% upside")
    return "  ·  ".join(parts)


def _add_key_metrics_table(doc: Document, metrics: list) -> None:
    """Render the cover's headline metrics as a two-column table:
    label on the left, value (with optional change + tone color) on
    the right. Two columns plays well in landscape Word docs and
    keeps every metric on its own row even with long labels."""
    table = doc.add_table(rows=len(metrics), cols=2)
    table.style = "Light Grid Accent 1"
    for row_idx, metric in enumerate(metrics):
        label_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(metric.label)
        label_run.bold = True
        label_run.font.size = Pt(10)

        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(metric.value)
        value_run.font.size = Pt(10)
        if metric.change:
            change_run = value_para.add_run(f"  {metric.change}")
            change_run.font.size = Pt(9)
            tone_color = _TONE_COLORS.get(metric.tone or "", None)
            if tone_color is not None:
                change_run.font.color.rgb = tone_color
                change_run.bold = True


def _format_eyebrow(report: ReportV3) -> str:
    parts = [report.template_id]
    if report.created_at is not None:
        parts.append(report.created_at.strftime("%b %d, %Y"))
    return " · ".join(parts)


# ---------------------------------------------------------------------------
# Sections (markdown -> docx)
# ---------------------------------------------------------------------------


def _add_section(
    *,
    doc: Document,
    section: ReportV3Section,
    md: MarkdownIt,
    display_index_by_source_id: dict[str, int],
    charts_by_id: dict[str, ReportV3Chart],
) -> None:
    # Section title at Heading 2 (cover title is Heading 1 implicitly).
    doc.add_heading(section.title, level=2)

    # Rewrite citation markers + extract chart references before parsing
    # markdown. We strip ``{{chart:id}}`` from the prose and embed the
    # corresponding chart(s) inline at the end of the section body.
    rewritten, referenced_chart_ids = _prepare_markdown(
        section.markdown, display_index_by_source_id
    )

    tokens = md.parse(rewritten)
    _emit_tokens(doc, tokens)

    for chart_id in referenced_chart_ids:
        chart = charts_by_id.get(chart_id)
        if chart is None:
            continue
        _embed_chart(doc, chart)


def _prepare_markdown(
    markdown: str, display_index_by_source_id: dict[str, int]
) -> tuple[str, list[str]]:
    """Strip ``{{chart:id}}`` placeholders (returns the ids in
    first-appearance order, deduplicated) and rewrite
    ``[^source_id]`` markers to ``[N]`` numbers."""
    import re

    chart_ids: list[str] = []
    seen_charts: set[str] = set()

    def _take_chart(m: re.Match[str]) -> str:
        cid = m.group(1)
        if cid not in seen_charts:
            seen_charts.add(cid)
            chart_ids.append(cid)
        return ""

    cleaned = strip_anthropic_citation_markup(markdown)
    stripped = re.sub(r"\{\{chart:([a-z0-9_]+)\}\}", _take_chart, cleaned)

    def _rewrite_citation(m: re.Match[str]) -> str:
        sid = m.group(1)
        n = display_index_by_source_id.get(sid)
        return f"[{n}]" if n is not None else f"[{sid}]"

    rewritten = re.sub(r"\[\^([a-z0-9_]+)\]", _rewrite_citation, stripped)
    return rewritten, chart_ids


def _emit_tokens(doc: Document, tokens: list[Any]) -> None:
    """Walk markdown-it tokens and emit docx paragraphs/runs.

    Implements a focused subset of CommonMark sufficient for v3
    section bodies: headings, paragraphs, bold / italic / code-inline /
    links, ordered + bullet lists, blockquotes, soft breaks. Tables
    are passed through best-effort (markdown-it table plugin emits
    th/td tokens we map to a docx ``add_table``).

    Unknown token types are silently skipped — better to drop an
    exotic block than crash the whole export.
    """
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            level = _BODY_HEADING_BASE_LEVEL + (int(tok.tag[1:]) - 1)
            level = min(level, _MAX_HEADING_LEVEL)
            # Find the matching inline token (always the next one).
            inline = tokens[i + 1]
            doc.add_heading(_inline_to_text(inline), level=level)
            # Skip past heading_open + inline + heading_close.
            i += 3
            continue
        if tok.type == "paragraph_open":
            inline = tokens[i + 1]
            para = doc.add_paragraph()
            _add_inline_runs(para, inline)
            i += 3
            continue
        if tok.type == "bullet_list_open":
            end_idx = _find_matching(tokens, i, "bullet_list_close")
            _emit_list(doc, tokens[i + 1 : end_idx], style="List Bullet")
            i = end_idx + 1
            continue
        if tok.type == "ordered_list_open":
            end_idx = _find_matching(tokens, i, "ordered_list_close")
            _emit_list(doc, tokens[i + 1 : end_idx], style="List Number")
            i = end_idx + 1
            continue
        if tok.type == "blockquote_open":
            end_idx = _find_matching(tokens, i, "blockquote_close")
            for inner in tokens[i + 1 : end_idx]:
                if inner.type == "inline":
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.25)
                    _add_inline_runs(para, inner)
            i = end_idx + 1
            continue
        if tok.type == "table_open":
            end_idx = _find_matching(tokens, i, "table_close")
            _emit_table(doc, tokens[i + 1 : end_idx])
            i = end_idx + 1
            continue
        if tok.type == "fence" or tok.type == "code_block":
            para = doc.add_paragraph()
            run = para.add_run(tok.content.rstrip("\n"))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            i += 1
            continue
        if tok.type == "hr":
            para = doc.add_paragraph("———")
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue
        # Anything else (html_block, etc.) — skip.
        i += 1


def _find_matching(tokens: list[Any], start: int, close_type: str) -> int:
    """Return the index of the matching close token for the block
    that opens at ``tokens[start]``. Uses markdown-it's ``level``
    on each token to handle nested blocks correctly."""
    base_level = tokens[start].level
    for j in range(start + 1, len(tokens)):
        if tokens[j].type == close_type and tokens[j].level == base_level:
            return j
    # Malformed token stream — return the end so we skip rather than crash.
    return len(tokens) - 1


def _emit_list(doc: Document, inner_tokens: list[Any], *, style: str) -> None:
    i = 0
    while i < len(inner_tokens):
        tok = inner_tokens[i]
        if tok.type == "list_item_open":
            end_idx = _find_matching(inner_tokens, i, "list_item_close")
            # First inline child holds the bullet text.
            for child in inner_tokens[i + 1 : end_idx]:
                if child.type == "inline":
                    para = doc.add_paragraph(style=style)
                    _add_inline_runs(para, child)
                    break
            i = end_idx + 1
            continue
        i += 1


def _emit_table(doc: Document, inner_tokens: list[Any]) -> None:
    rows: list[list[str]] = []
    current_row: list[str] | None = None
    capture_inline = False
    for tok in inner_tokens:
        if tok.type in {"tr_open"}:
            current_row = []
        elif tok.type in {"th_open", "td_open"}:
            capture_inline = True
        elif tok.type == "inline" and capture_inline:
            current_row.append(_inline_to_text(tok)) if current_row is not None else None
            capture_inline = False
        elif tok.type in {"tr_close"} and current_row is not None:
            rows.append(current_row)
            current_row = None
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = row[c_idx] if c_idx < len(row) else ""


def _add_inline_runs(para, inline) -> None:
    """Walk the children of an inline token and add appropriately
    styled runs to the paragraph."""
    bold = False
    italic = False
    code = False
    pending_link_href: str | None = None
    for child in inline.children or []:
        t = child.type
        if t == "strong_open":
            bold = True
        elif t == "strong_close":
            bold = False
        elif t == "em_open":
            italic = True
        elif t == "em_close":
            italic = False
        elif t == "code_inline":
            run = para.add_run(child.content)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif t == "link_open":
            pending_link_href = child.attrGet("href") or ""
        elif t == "link_close":
            pending_link_href = None
        elif t == "softbreak":
            para.add_run(" ")
        elif t == "hardbreak":
            run = para.add_run()
            run.add_break()
        elif t == "text":
            if pending_link_href:
                _add_hyperlink(para, child.content, pending_link_href)
            else:
                run = para.add_run(child.content)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if code:
                    run.font.name = "Courier New"


def _inline_to_text(inline) -> str:
    """Flatten an inline token's text content for headings + table
    cells (no run-level styling needed there)."""
    out: list[str] = []
    for child in inline.children or []:
        if child.type == "text":
            out.append(child.content)
        elif child.type == "code_inline":
            out.append(child.content)
        elif child.type == "softbreak" or child.type == "hardbreak":
            out.append(" ")
    return "".join(out)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a clickable hyperlink to ``paragraph``. python-docx 1.x
    has no high-level hyperlink API, so build the relationship + the
    underlying XML by hand. Falls back to plain text on errors."""
    try:
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
        new_run = paragraph._p.makeelement(qn("w:r"), {})
        rPr = paragraph._p.makeelement(qn("w:rPr"), {})
        u = paragraph._p.makeelement(qn("w:u"), {qn("w:val"): "single"})
        color = paragraph._p.makeelement(qn("w:color"), {qn("w:val"): "1F6FEB"})
        rPr.append(u)
        rPr.append(color)
        new_run.append(rPr)
        t = paragraph._p.makeelement(qn("w:t"), {})
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:  # pragma: no cover
        log.exception("hyperlink emit failed; falling back to plain text")
        paragraph.add_run(text)


# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------


def _embed_chart(doc: Document, chart: ReportV3Chart) -> None:
    """Render the chart via matplotlib and embed the PNG. Falls back
    to a typed paragraph placeholder when rendering fails (no
    matplotlib in env, unsupported chart type, etc.)."""
    title = doc.add_paragraph()
    title_run = title.add_run(chart.title)
    title_run.bold = True
    title_run.font.size = Pt(10)

    png = _try_render_chart_png(chart)
    if png is not None:
        try:
            doc.add_picture(io.BytesIO(png), width=Inches(5.5))
            return
        except Exception:  # pragma: no cover
            log.exception("chart %s failed to embed; emitting placeholder", chart.chart_id)

    placeholder = doc.add_paragraph()
    placeholder_run = placeholder.add_run(f"[chart placeholder: {chart.chart_type}]")
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _try_render_chart_png(chart: ReportV3Chart) -> bytes | None:
    """Round-trip the persisted spec_json through ``ChartSpec`` and
    the v3 chart renderer. Returns None when matplotlib is absent or
    the spec can't be rendered."""
    try:
        spec = ChartSpec.model_validate_json(chart.spec_json)
    except Exception:
        log.warning("chart %s has invalid spec_json", chart.chart_id)
        return None
    try:
        from openlia.llm.runtime.report_v3.rendering.chart_renderer import (
            render_chart_png,
        )
    except ImportError:  # pragma: no cover
        return None
    try:
        rendered = render_chart_png(spec)
    except Exception:
        log.exception("chart_renderer crashed for %s", chart.chart_id)
        return None
    if rendered is None:
        return None
    return rendered.png_bytes


# ---------------------------------------------------------------------------
# Bibliography
# ---------------------------------------------------------------------------


def _add_bibliography(doc: Document, citations: list[ReportV3Citation]) -> None:
    """Render every persisted citation. Body-linked rows keep the
    ``display_index`` the renderer assigned; the rest get sequential
    numbers appended after the highest assigned one so the reader
    always sees every source consulted, even when marker linkage
    misfires (Anthropic native search emits non-source-id markers we
    can't link back automatically)."""
    if not citations:
        return
    indexed: list[tuple[int, ReportV3Citation]] = []
    unindexed: list[ReportV3Citation] = []
    max_index = 0
    for c in citations:
        if c.display_index is None:
            unindexed.append(c)
            continue
        max_index = max(max_index, c.display_index)
        indexed.append((c.display_index, c))
    next_index = max_index + 1
    for c in unindexed:
        indexed.append((next_index, c))
        next_index += 1
    indexed.sort(key=lambda pair: pair[0])
    doc.add_heading("Sources", level=2)
    for idx, c in indexed:
        prov = _parse_provenance(c.provenance_json)
        url = prov.get("url") if isinstance(prov.get("url"), str) else None
        title = prov.get("title") if isinstance(prov.get("title"), str) else None
        para = doc.add_paragraph()
        idx_run = para.add_run(f"[{idx}] ")
        idx_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        if url:
            _add_hyperlink(para, title or url, url)
        else:
            para.add_run(title or c.tool_name)
        suffix = para.add_run(f"  ({c.source_id})")
        suffix.font.size = Pt(9)
        suffix.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)


def _parse_provenance(raw: str) -> dict[str, Any]:
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else {}
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}


__all__ = ["render_docx"]
