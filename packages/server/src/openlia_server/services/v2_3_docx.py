"""DOCX renderer for v2.3 equity-research reports.

Consumes a completed ``ReportState`` (with ``state.resolved`` populated by
AssembleStage) and returns ``.docx`` bytes. Sections are rendered in
outline order; each section's resolved body keeps the deterministic
``[^N]`` footnote markers from ``schemas.resolve()``, then those markers
get rewritten into native Word footnote references and the footnotes
themselves into a real ``word/footnotes.xml`` part (see
``v2_3_docx_native``).

Charts try a matplotlib-rendered PNG embedded via ``add_picture``
first (``v2_3_docx_chart``); when rendering is impossible (no
matplotlib, mis-shaped data, unsupported chart type) the renderer
falls back to a data table so the figure is never silently dropped.
The chart-data table view is still emitted alongside the image as a
machine-readable backup.

A "References" section is still rendered as a tail for readers /
viewers that don't honor native footnotes (notably some web preview
tools). Native + tail dual-render is intentional belt-and-suspenders.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from openlia.llm.runtime.report_v2_3.schemas import (
    BundleFact,
    BundleSeries,
    ChartSpec,
    ReportThesis,
    ResearchBundle,
    ResolvedReport,
)
from openlia.llm.runtime.report_v2_3.state import ReportState

from .v2_3_docx_chart import try_render_chart_png
from .v2_3_docx_native import attach_native_footnotes


def render_docx(state: ReportState) -> bytes:
    """Render a v2.3 ReportState into a Word document.

    Requires ``state.resolved`` to be populated (i.e. AssembleStage has
    successfully run). Raises ``RuntimeError`` otherwise — callers can
    detect a not-yet-complete run by checking ``state.resolved`` before
    calling.
    """
    if state.resolved is None:
        raise RuntimeError("render_docx requires state.resolved to be set by ASSEMBLE.")
    if state.outline is None or state.thesis is None or state.bundle is None:
        raise RuntimeError("render_docx requires outline + thesis + bundle on state.")

    doc = Document()
    _configure_default_styles(doc)
    _add_cover(doc, state)
    _add_thesis_block(doc, state.thesis)

    charts_by_section = _index_charts(state.thesis.charts)

    sections_by_id = {s.section_id: s for s in state.sections}
    for outline_section in state.outline.sections:
        body = state.resolved.section_bodies.get(outline_section.id)
        if body is None:
            continue
        written = sections_by_id.get(outline_section.id)
        title = written.title if written is not None else outline_section.title

        doc.add_heading(title, level=1)
        for paragraph_text in body.split("\n\n"):
            stripped = paragraph_text.strip()
            if stripped:
                doc.add_paragraph(stripped)

        for chart in charts_by_section.get(outline_section.id, []):
            _add_chart_data_table(doc, chart, state.bundle, state.resolved.figure_labels)

    _add_references(doc, state.resolved)
    # Rewrite the deterministic ``[^N]`` markers AssembleStage produced into
    # real Word footnote references + a footnotes.xml package part. Done
    # AFTER all paragraphs are in the doc so the rewrite covers every
    # body paragraph in one pass.
    attach_native_footnotes(doc, state.resolved.footnotes)
    return _serialize(doc)


# ---------------------------------------------------------------------------
# Cover + thesis block
# ---------------------------------------------------------------------------


def _configure_default_styles(doc: Any) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_cover(doc: Any, state: ReportState) -> None:
    cover_title = ", ".join(state.tickers)
    h = doc.add_heading(cover_title, level=0)
    for run in h.runs:
        run.font.size = Pt(28)

    if state.thesis is not None:
        subtitle = state.thesis.central_argument
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Report type: {state.report_type.value}  ·  Language: {state.language.value}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_thesis_block(doc: Any, thesis: ReportThesis) -> None:
    doc.add_heading("Thesis", level=1)
    doc.add_paragraph(thesis.central_argument)

    if thesis.key_takeaways:
        doc.add_heading("Key takeaways", level=2)
        for takeaway in thesis.key_takeaways:
            doc.add_paragraph(takeaway, style="List Bullet")

    if thesis.valuation_stance:
        doc.add_heading("Valuation stance", level=2)
        doc.add_paragraph(thesis.valuation_stance)


# ---------------------------------------------------------------------------
# Chart data table
# ---------------------------------------------------------------------------


def _index_charts(charts: list[ChartSpec]) -> dict[str, list[ChartSpec]]:
    out: dict[str, list[ChartSpec]] = {}
    for chart in charts:
        out.setdefault(chart.section_id, []).append(chart)
    return out


def _add_chart_data_table(
    doc: Any,
    chart: ChartSpec,
    bundle: ResearchBundle,
    figure_labels: dict[str, int],
) -> None:
    figure_n = figure_labels.get(chart.id)
    caption_text = f"Figure {figure_n}: {chart.title}" if figure_n is not None else chart.title
    caption = doc.add_paragraph()
    caption_run = caption.add_run(caption_text)
    caption_run.bold = True

    if chart.claim:
        claim = doc.add_paragraph(chart.claim)
        for run in claim.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Try a real graphic first. The data table follows as a backup so
    # screen readers / Word search still see the numbers.
    png = try_render_chart_png(chart, bundle)
    if png is not None:
        import io as _io

        doc.add_picture(_io.BytesIO(png), width=Inches(6.0))

    table = doc.add_table(rows=1 + len(chart.category_labels), cols=1 + len(chart.series))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"

    # Header row: category-axis label + one column per series.
    header_cells = table.rows[0].cells
    header_cells[0].text = chart.x_axis_label or "Category"
    for col_idx, series in enumerate(chart.series, start=1):
        header_cells[col_idx].text = series.name

    # Body: one row per category. Series cells are resolved by:
    #   - if series has exactly one fact_id and that fact is a BundleSeries
    #     of matching length, pull the per-period values
    #   - otherwise, the cell is the scalar value of the fact_id at the
    #     same index (or blank if out of range / missing)
    for row_idx, category in enumerate(chart.category_labels, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = category
        for col_idx, series in enumerate(chart.series, start=1):
            cells[col_idx].text = _series_cell_value(series, row_idx - 1, bundle)


def _series_cell_value(series: Any, idx: int, bundle: ResearchBundle) -> str:
    if len(series.value_fact_ids) == 1:
        fact = bundle.facts.get(series.value_fact_ids[0])
        if fact is not None and isinstance(fact.value, BundleSeries):
            if idx < len(fact.value.points):
                return _fmt(fact.value.points[idx].value)
            return ""

    if idx < len(series.value_fact_ids):
        fact = bundle.facts.get(series.value_fact_ids[idx])
        if fact is None:
            return ""
        return _fmt_fact_scalar(fact)
    return ""


def _fmt_fact_scalar(fact: BundleFact) -> str:
    if isinstance(fact.value, (int, float)):
        return _fmt(fact.value)
    if isinstance(fact.value, BundleSeries):
        # The series shape was already handled above; if we got here the
        # category labels don't line up, so fall back to a compact summary.
        return f"{len(fact.value.points)} pts"
    return str(fact.value)


def _fmt(value: float) -> str:
    if abs(value) >= 1000:
        return f"{value:,.0f}"
    if abs(value) >= 1:
        return f"{value:,.2f}"
    return f"{value:.3f}"


# ---------------------------------------------------------------------------
# References
# ---------------------------------------------------------------------------


def _add_references(doc: Any, resolved: ResolvedReport) -> None:
    if not resolved.footnotes:
        return
    doc.add_page_break()
    doc.add_heading("References", level=1)
    for idx, footnote in enumerate(resolved.footnotes, start=1):
        doc.add_paragraph(f"[{idx}] {footnote}")


# ---------------------------------------------------------------------------
# Serialization
# ---------------------------------------------------------------------------


def _serialize(doc: Any) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["render_docx"]
