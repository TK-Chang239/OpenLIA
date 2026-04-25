from __future__ import annotations

import asyncio
import io
from typing import Any

from playwright.async_api import Browser, Playwright, async_playwright


class BrowserLauncher:
    def __init__(self) -> None:
        self._lock = asyncio.Lock()
        self._playwright: Playwright | None = None
        self._browser: Browser | None = None
        self._closed = False

    async def browser(self) -> Browser:
        async with self._lock:
            if self._closed:
                raise RuntimeError("BrowserLauncher has been shut down")
            if self._browser is None:
                self._playwright = await async_playwright().start()
                self._browser = await self._playwright.chromium.launch(headless=True)
            return self._browser

    async def shutdown(self) -> None:
        async with self._lock:
            if self._closed:
                return
            self._closed = True
            if self._browser is not None:
                await self._browser.close()
                self._browser = None
            if self._playwright is not None:
                await self._playwright.stop()
                self._playwright = None


def export_report_docx(schema: dict[str, Any]) -> bytes:
    """Render a report schema dict to .docx bytes.

    Maps the canonical schema (cover + sections + blocks) to a flat Word
    document. Field names align with Phase 13 P1-01:
        - metric_cards.metrics
        - table.headers, table.rows
        - key_finding.content
        - rating_badge.rating
    Charts are rendered as a placeholder paragraph (`[Chart: <title>]`)
    since Word has no native ECharts equivalent.
    """
    from docx import Document  # local import — server-only optional dep
    from docx.shared import Pt

    doc = Document()
    cover = schema.get("cover") or {}
    title = str(cover.get("title", "Report"))
    doc.add_heading(title, level=0)
    if cover.get("subtitle"):
        doc.add_paragraph(str(cover["subtitle"]))
    if cover.get("ticker"):
        p = doc.add_paragraph()
        run = p.add_run(str(cover["ticker"]))
        run.bold = True
    if cover.get("tagline"):
        doc.add_paragraph(str(cover["tagline"]))
    key_metrics = cover.get("key_metrics") or []
    for m in key_metrics:
        label = str(m.get("label", ""))
        value = str(m.get("value", ""))
        delta = m.get("delta")
        line = f"{label}: {value}" + (f"  ({delta})" if delta else "")
        doc.add_paragraph(line)

    for section in schema.get("sections", []) or []:
        doc.add_heading(str(section.get("title", "")), level=1)
        for block in section.get("blocks", []) or []:
            _render_docx_block(doc, block)

    furniture = schema.get("page_furniture") or {}
    if isinstance(furniture, dict):
        disclaimer = furniture.get("disclaimer")
        if disclaimer:
            p = doc.add_paragraph()
            run = p.add_run(str(disclaimer))
            run.italic = True
            run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_block(doc: Any, block: dict[str, Any]) -> None:
    btype = str(block.get("type", ""))
    if btype == "text":
        doc.add_paragraph(str(block.get("content", "")))
        return
    if btype == "key_finding":
        p = doc.add_paragraph()
        run = p.add_run(str(block.get("content", "")))
        run.bold = True
        return
    if btype == "rating_badge":
        rating = str(block.get("rating", ""))
        prev = block.get("previous_rating")
        suffix = f" (prev: {prev})" if prev else ""
        p = doc.add_paragraph()
        run = p.add_run(f"Rating: {rating}{suffix}")
        run.bold = True
        return
    if btype == "metric_cards":
        metrics = block.get("metrics", []) or []
        for m in metrics:
            label = str(m.get("label", ""))
            value = str(m.get("value", ""))
            delta = m.get("delta")
            line = f"{label}: {value}" + (f"  ({delta})" if delta else "")
            doc.add_paragraph(line, style="List Bullet")
        return
    if btype == "table":
        title = block.get("title")
        if title:
            p = doc.add_paragraph()
            run = p.add_run(str(title))
            run.bold = True
        headers = block.get("headers", []) or []
        rows = block.get("rows", []) or []
        if not headers:
            return
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h.get("label", ""))
        for row in rows:
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                key = h.get("key", "")
                val = row.get(key, "")
                row_cells[i].text = "" if val is None else str(val)
        for fn in block.get("footnotes", []) or []:
            p = doc.add_paragraph()
            run = p.add_run(str(fn))
            run.italic = True
        return
    if btype == "group":
        for child in block.get("blocks", []) or []:
            _render_docx_block(doc, child)
        return
    # Charts and unknown block types: placeholder paragraph.
    title = block.get("title") or btype.replace("_", " ").title()
    p = doc.add_paragraph()
    run = p.add_run(f"[{title}]")
    run.italic = True


async def export_report_pdf(
    launcher: BrowserLauncher,
    html: str,
    *,
    header_html: str | None = None,
    footer_html: str | None = None,
    bundle_url: str | None = None,
    cookies: list[dict[str, Any]] | None = None,
) -> bytes:
    """Render a report HTML body to PDF bytes via Playwright.

    Two paths:
      - When `bundle_url` is provided: navigate the page to that URL with
        `wait_until="networkidle"`. Use this for the SPA-driven flow where
        the React `ReportRenderer` mounts and ECharts renders real vector
        graphics. Optional `cookies` are forwarded so the SPA can call
        `/api/reports/:id` against the protected backend.
      - Otherwise: use `page.set_content(html)`. This is the static-fallback
        path that ships chart titles, tables, and metric cards baked into
        the HTML — used by tests and any environment where the SPA bundle
        isn't reachable.
    """
    browser = await launcher.browser()
    context = await browser.new_context()
    try:
        if cookies:
            await context.add_cookies(cookies)
        page = await context.new_page()
        if bundle_url is not None:
            await page.goto(bundle_url, wait_until="networkidle")
        else:
            await page.set_content(html, wait_until="networkidle")
        kwargs: dict[str, Any] = {
            "format": "A4",
            "margin": {"top": "20mm", "bottom": "25mm", "left": "20mm", "right": "20mm"},
            "print_background": True,
        }
        if header_html or footer_html:
            kwargs["display_header_footer"] = True
            if header_html:
                kwargs["header_template"] = header_html
            if footer_html:
                kwargs["footer_template"] = footer_html
        return await page.pdf(**kwargs)
    finally:
        await context.close()
