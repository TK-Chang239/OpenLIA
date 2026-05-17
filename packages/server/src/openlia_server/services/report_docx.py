"""High-fidelity DOCX export for OpenLIA reports.

Walks the canonical ReportSchema and emits a Word-native document with:
- A styled cover page (title, subtitle, ticker, tagline, key-metric table).
- A Word native TOC field at the top (user refreshes via right-click).
- One Heading 1 per section.
- Native Word tables for `table` and `metric_cards` blocks.
- Bullet lists, italicized quotes, callout grids, timeline rows, etc.
- **Chart blocks embed pre-screenshotted PNG bytes** (from
  `capture_chart_pngs`) at 6.5"-wide so the visual fidelity matches the
  browser print page.
- Native Word header/footer carrying the report title and page X of Y.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_CHART_TYPES = {
    "line_chart",
    "bar_chart",
    "area_chart",
    "pie_chart",
    "candlestick_chart",
    "waterfall_chart",
    "scatter_plot",
    "heatmap",
    "treemap",
    "combo_chart",
}


def assemble_docx(
    schema: dict[str, Any],
    *,
    chart_pngs: dict[str, bytes],
    header_text: str = "",
) -> bytes:
    doc = Document()
    _configure_default_styles(doc)
    _set_header_footer(doc, header_text=header_text, schema=schema)
    _add_cover(doc, schema.get("cover") or {})
    _add_toc_field(doc)

    sections = schema.get("sections") or []
    for sec_idx, section in enumerate(sections):
        if sec_idx > 0:
            doc.add_page_break()
        doc.add_heading(str(section.get("title", "")), level=1)
        for blk_idx, block in enumerate(section.get("blocks") or []):
            _render_block(
                doc,
                block,
                path=f"{sec_idx}-{blk_idx}",
                chart_pngs=chart_pngs,
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _configure_default_styles(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_cover(doc, cover: dict[str, Any]) -> None:
    title = str(cover.get("title", "Report"))
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.size = Pt(28)
    subtitle = cover.get("subtitle")
    if subtitle:
        p = doc.add_paragraph(str(subtitle))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    ticker = cover.get("ticker")
    if ticker:
        p = doc.add_paragraph()
        run = p.add_run(str(ticker))
        run.bold = True
    tagline = cover.get("tagline")
    if tagline:
        doc.add_paragraph(str(tagline))
    metrics = cover.get("key_metrics") or []
    if metrics:
        tbl = doc.add_table(rows=len(metrics), cols=2)
        try:
            tbl.style = "Light Grid Accent 1"
        except KeyError:
            tbl.style = "Table Grid"
        for i, m in enumerate(metrics):
            tbl.rows[i].cells[0].text = str(m.get("label", ""))
            value = str(m.get("value", ""))
            delta = m.get("delta")
            if delta:
                value = f"{value}  ({delta})"
            tbl.rows[i].cells[1].text = value
    doc.add_page_break()


def _add_toc_field(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click to update Table of Contents"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)
    doc.add_page_break()


def _set_header_footer(doc, *, header_text: str, schema: dict[str, Any]) -> None:
    section = doc.sections[0]
    if header_text:
        section.header.paragraphs[0].text = header_text
    footer_p = section.footer.paragraphs[0]
    furniture = (schema.get("page_furniture") or {}).get("footer") or {}
    subject = str(furniture.get("subject") or furniture.get("left") or "")
    date = str(furniture.get("date") or "")
    left = " · ".join(s for s in [subject, date] if s)
    footer_p.text = left
    footer_p.add_run("\t\tPage ")
    _append_field(footer_p.add_run(), "PAGE")
    footer_p.add_run(" of ")
    _append_field(footer_p.add_run(), "NUMPAGES")


def _append_field(run, instr: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run._r.append(fld)


def _render_block(
    doc,
    block: dict[str, Any],
    *,
    path: str,
    chart_pngs: dict[str, bytes],
) -> None:
    btype = str(block.get("type", ""))

    if btype in _CHART_TYPES:
        title = block.get("title")
        if title:
            p = doc.add_paragraph()
            run = p.add_run(str(title))
            run.bold = True
        png = chart_pngs.get(path)
        if png:
            doc.add_picture(io.BytesIO(png), width=Inches(6.5))
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"[chart: {title or btype}]")
            run.italic = True
        return

    if btype == "text":
        doc.add_paragraph(str(block.get("content", "")))
        return

    if btype == "key_finding":
        p = doc.add_paragraph()
        run = p.add_run(str(block.get("content", "")))
        run.bold = True
        return

    if btype == "rating_badge":
        rating = str(block.get("rating", ""))
        prev = block.get("previous_rating")
        suffix = f" (prev: {prev})" if prev else ""
        p = doc.add_paragraph()
        run = p.add_run(f"Rating: {rating}{suffix}")
        run.bold = True
        return

    if btype == "metric_cards":
        metrics = block.get("metrics") or []
        if metrics:
            tbl = doc.add_table(rows=len(metrics), cols=2)
            try:
                tbl.style = "Light Grid Accent 1"
            except KeyError:
                tbl.style = "Table Grid"
            for i, m in enumerate(metrics):
                tbl.rows[i].cells[0].text = str(m.get("label", ""))
                value = str(m.get("value", ""))
                delta = m.get("delta")
                if delta:
                    value = f"{value}  ({delta})"
                tbl.rows[i].cells[1].text = value
        return

    if btype == "table":
        title = block.get("title")
        if title:
            p = doc.add_paragraph()
            run = p.add_run(str(title))
            run.bold = True
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        if not headers:
            return
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h.get("label", ""))
        for row in rows:
            row_cells = tbl.add_row().cells
            for i, h in enumerate(headers):
                key = h.get("key", "")
                val = row.get(key, "")
                row_cells[i].text = "" if val is None else str(val)
        for fn in block.get("footnotes") or []:
            p = doc.add_paragraph()
            run = p.add_run(str(fn))
            run.italic = True
        return

    if btype == "bullet_list":
        for item in block.get("items") or []:
            doc.add_paragraph(str(item), style="List Bullet")
        return

    if btype in ("pull_quote", "quote"):
        content = str(block.get("content") or block.get("text") or "")
        try:
            doc.add_paragraph(content, style="Intense Quote")
        except KeyError:
            p = doc.add_paragraph(content)
            for run in p.runs:
                run.italic = True
        attribution = block.get("attribution") or block.get("source")
        if attribution:
            p2 = doc.add_paragraph(f"— {attribution}")
            for run in p2.runs:
                run.italic = True
        return

    if btype == "callout_grid":
        items = block.get("items") or []
        if items:
            cols = 2
            rows = (len(items) + cols - 1) // cols
            tbl = doc.add_table(rows=rows, cols=cols)
            try:
                tbl.style = "Light Shading"
            except KeyError:
                tbl.style = "Table Grid"
            for idx, item in enumerate(items):
                r, c = divmod(idx, cols)
                cell = tbl.rows[r].cells[c]
                head = str(item.get("title") or "")
                body = str(item.get("body") or item.get("content") or "")
                cell.text = f"{head}\n{body}" if head else body
        return

    if btype == "timeline":
        for event in block.get("events") or []:
            date = str(event.get("date") or "")
            content = str(event.get("content") or event.get("title") or "")
            doc.add_paragraph(f"{date}: {content}", style="List Bullet")
        return

    if btype == "comparison_split":
        cols = block.get("columns") or []
        if len(cols) == 2:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            for i, col in enumerate(cols):
                head = str(col.get("title") or "")
                body = str(col.get("body") or col.get("content") or "")
                tbl.rows[0].cells[i].text = f"{head}\n\n{body}" if head else body
        return

    if btype == "group":
        for child_idx, child in enumerate(block.get("blocks") or []):
            _render_block(doc, child, path=f"{path}.{child_idx}", chart_pngs=chart_pngs)
        return

    # Fallback for unknown block types: italic placeholder.
    p = doc.add_paragraph()
    run = p.add_run(f"[{btype}]")
    run.italic = True
